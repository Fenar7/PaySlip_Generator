from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "phase1-qa-test-plan.docx"


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1D, 0x17, 0x10)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x1D, 0x17, 0x10)


def configure_section(section) -> None:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.text = "Page "
    add_page_number(footer_p)


def add_title_page(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(10)

    document.add_paragraph()
    document.add_paragraph()

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Business Document Generator")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Phase 1 QA Test Plan")
    run.bold = True
    run.font.size = Pt(16)

    document.add_paragraph()

    for label, value in [
        ("Target Environment", "Vercel Preview"),
        ("Version", "v1.0"),
        ("Prepared For", "Testing Team"),
        ("Date", date.today().strftime("%d %B %Y")),
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(f"{label}: ").bold = True
        p.add_run(value)

    document.add_page_break()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def add_test_case_table(document: Document, cases: list[dict[str, str]]) -> None:
    headers = [
        "Test Case ID",
        "Module",
        "Scenario",
        "Preconditions",
        "Steps",
        "Expected Result",
        "Priority",
    ]
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [0.9, 0.85, 1.25, 1.35, 2.35, 2.25, 0.7]
    for cell, header, width in zip(table.rows[0].cells, headers, widths):
        cell.width = Inches(width)
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "EADFCB")

    for case in cases:
        row = table.add_row().cells
        for index, key in enumerate(
            ["id", "module", "scenario", "preconditions", "steps", "expected", "priority"]
        ):
            row[index].width = Inches(widths[index])
            set_cell_text(row[index], case[key])

    document.add_paragraph()


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])

    add_title_page(document)

    second = document.sections[-1]
    second.start_type = WD_SECTION.NEW_PAGE
    configure_section(second)

    document.add_heading("Document Purpose", level=1)
    document.add_paragraph(
        "This document provides the testing team with a structured Phase 1 QA handoff for the Business Document Generator. It defines the product scope, test environment assumptions, detailed manual test cases, and the information required for defect logging and final sign-off."
    )
    document.add_paragraph(
        "QA should use this pack to validate the deployed Vercel Preview build of the voucher, salary slip, and invoice generators, including live preview behavior, print flow, PDF export, PNG export, and responsive usability."
    )

    document.add_heading("In Scope", level=1)
    add_bullets(
        document,
        [
            "Voucher Generator",
            "Salary Slip Generator",
            "Invoice Generator",
            "Live preview updates",
            "Template switching",
            "Branding controls including logo upload and accent color",
            "Print flow",
            "PDF export",
            "PNG export",
            "Desktop, tablet, and mobile usability",
            "Serverless behavior on Vercel Preview",
        ],
    )

    document.add_heading("Out of Scope", level=1)
    add_bullets(
        document,
        [
            "Authentication and user accounts",
            "Saved drafts or document history",
            "Cloud persistence",
            "Team collaboration and approval workflows",
            "Recurring automation",
            "Database-backed functionality",
        ],
    )

    document.add_heading("Test Environment and Preconditions", level=1)
    add_bullets(
        document,
        [
            "Primary target: Vercel Preview deployment",
            "Preview URL: ______________________________",
            "Preferred browser: Chromium / Google Chrome",
            "Stable network connectivity is required",
            "Allow browser pop-ups for print validation",
            "Allow file downloads for PDF and PNG checks",
            "Use at least one desktop viewport and one smaller tablet/mobile viewport during the run",
        ],
    )

    document.add_heading("Entry and Exit Criteria", level=1)
    document.add_heading("Entry Criteria", level=2)
    add_bullets(
        document,
        [
            "The Vercel Preview deployment URL is reachable.",
            "The `/`, `/voucher`, `/salary-slip`, and `/invoice` routes load successfully.",
            "Export endpoints respond without deployment errors.",
        ],
    )
    document.add_heading("Exit Criteria", level=2)
    add_bullets(
        document,
        [
            "All Blocker and High-priority test cases are passed.",
            "No unresolved defect blocks voucher, salary slip, or invoice creation or export.",
            "Print, PDF, and PNG flows complete successfully for each module.",
        ],
    )

    document.add_heading("Severity / Priority Guide", level=1)
    add_bullets(
        document,
        [
            "Blocker: Core route or export flow is unusable.",
            "High: Major workflow issue with incorrect output, broken layout, or failed export.",
            "Medium: Functional issue with workaround available.",
            "Low: Cosmetic or minor usability issue that does not block completion.",
        ],
    )

    document.add_heading("Module Test Cases", level=1)

    home_cases = [
        {
            "id": "NAV-001",
            "module": "Home",
            "scenario": "Landing page loads",
            "preconditions": "Preview deployment is reachable.",
            "steps": "Open the preview URL. Wait for the home page to finish loading.",
            "expected": "The landing page loads without error and shows entry points for voucher, salary slip, and invoice modules.",
            "priority": "High",
        },
        {
            "id": "NAV-002",
            "module": "Navigation",
            "scenario": "Workspace links open correctly",
            "preconditions": "Home page is open.",
            "steps": "Open the voucher workspace. Return to home. Open the salary slip workspace. Return to home. Open the invoice workspace.",
            "expected": "Each workspace route loads correctly and shows the correct generator heading and controls.",
            "priority": "High",
        },
    ]

    voucher_cases = [
        {
            "id": "VOU-001",
            "module": "Voucher",
            "scenario": "Voucher workspace renders with defaults",
            "preconditions": "Open `/voucher`.",
            "steps": "Wait for the workspace to load and inspect the form and preview.",
            "expected": "The voucher form, preview, template selector, and export buttons are visible with editable default values.",
            "priority": "High",
        },
        {
            "id": "VOU-002",
            "module": "Voucher",
            "scenario": "Voucher type updates preview content",
            "preconditions": "Voucher workspace is open.",
            "steps": "Change voucher type from payment to receipt.",
            "expected": "The preview updates to receipt wording and the counterparty label changes accordingly.",
            "priority": "High",
        },
        {
            "id": "VOU-003",
            "module": "Voucher",
            "scenario": "Template switching works",
            "preconditions": "Voucher workspace is open.",
            "steps": "Switch between Minimal Office and Traditional Ledger templates.",
            "expected": "The preview updates to the selected template without losing the entered form data.",
            "priority": "High",
        },
        {
            "id": "VOU-004",
            "module": "Voucher",
            "scenario": "Visibility toggle hides notes",
            "preconditions": "Voucher workspace is open with default notes visible.",
            "steps": "Turn off the Notes visibility toggle.",
            "expected": "The notes section disappears from the live preview and layout reflows cleanly.",
            "priority": "Medium",
        },
        {
            "id": "VOU-005",
            "module": "Voucher",
            "scenario": "PDF export completes",
            "preconditions": "Voucher workspace is open with valid data.",
            "steps": "Click Export PDF and download the file. Open the downloaded PDF.",
            "expected": "A PDF file downloads successfully, opens without corruption, and contains the expected voucher text.",
            "priority": "Blocker",
        },
        {
            "id": "VOU-006",
            "module": "Voucher",
            "scenario": "PNG export completes",
            "preconditions": "Voucher workspace is open with valid data.",
            "steps": "Click Export PNG and open the downloaded file.",
            "expected": "A PNG file downloads successfully and the document image is complete and readable.",
            "priority": "Blocker",
        },
    ]

    salary_cases = [
        {
            "id": "SAL-001",
            "module": "Salary Slip",
            "scenario": "Salary slip workspace renders with defaults",
            "preconditions": "Open `/salary-slip`.",
            "steps": "Wait for the workspace to load and inspect the form and preview.",
            "expected": "The salary slip form, preview, template selector, and export actions are visible with editable defaults.",
            "priority": "High",
        },
        {
            "id": "SAL-002",
            "module": "Salary Slip",
            "scenario": "Earnings and deductions update totals",
            "preconditions": "Salary slip workspace is open.",
            "steps": "Add a new earning row. Enter a label and amount. Review the preview totals.",
            "expected": "The new row appears in the preview and the total earnings and net salary update immediately.",
            "priority": "High",
        },
        {
            "id": "SAL-003",
            "module": "Salary Slip",
            "scenario": "Visibility toggles hide optional sections",
            "preconditions": "Salary slip workspace is open.",
            "steps": "Disable bank details, notes, and signature-related toggles one by one.",
            "expected": "The corresponding preview sections disappear and surrounding layout remains balanced.",
            "priority": "High",
        },
        {
            "id": "SAL-004",
            "module": "Salary Slip",
            "scenario": "Template switching works",
            "preconditions": "Salary slip workspace is open.",
            "steps": "Switch between Corporate Clean and Modern Premium templates.",
            "expected": "The selected template renders correctly and all entered values remain intact.",
            "priority": "High",
        },
        {
            "id": "SAL-005",
            "module": "Salary Slip",
            "scenario": "Print route renders correctly",
            "preconditions": "Salary slip workspace is open with valid data.",
            "steps": "Click Print Salary Slip and allow the print surface to open.",
            "expected": "The print page opens successfully and shows the correct salary-slip render-ready surface.",
            "priority": "Blocker",
        },
        {
            "id": "SAL-006",
            "module": "Salary Slip",
            "scenario": "PDF export completes",
            "preconditions": "Salary slip workspace is open with valid data.",
            "steps": "Click Export PDF and open the downloaded file.",
            "expected": "A salary-slip PDF downloads successfully and contains the expected employee and pay-period text.",
            "priority": "Blocker",
        },
        {
            "id": "SAL-007",
            "module": "Salary Slip",
            "scenario": "PNG export completes",
            "preconditions": "Salary slip workspace is open with valid data.",
            "steps": "Click Export PNG and open the downloaded file.",
            "expected": "A salary-slip PNG downloads successfully and the document image is complete and readable.",
            "priority": "Blocker",
        },
        {
            "id": "SAL-008",
            "module": "Salary Slip",
            "scenario": "PDF page splitting does not cut boxed sections",
            "preconditions": "Salary slip workspace is open with enough visible sections to extend the output vertically.",
            "steps": "Export the salary slip as PDF. Review the page boundary around bank details, notes, and signature sections.",
            "expected": "Bordered sections are not cut in the middle. If a section does not fit, it moves fully to the next page.",
            "priority": "High",
        },
    ]

    invoice_cases = [
        {
            "id": "INV-001",
            "module": "Invoice",
            "scenario": "Invoice workspace renders with defaults",
            "preconditions": "Open `/invoice`.",
            "steps": "Wait for the workspace to load and inspect the form and preview.",
            "expected": "The invoice form, preview, template selector, and export actions are visible with editable defaults.",
            "priority": "High",
        },
        {
            "id": "INV-002",
            "module": "Invoice",
            "scenario": "Line item math updates preview totals",
            "preconditions": "Invoice workspace is open.",
            "steps": "Change a discount amount and amount paid value. Review totals in the preview.",
            "expected": "Subtotal, tax, grand total, amount paid, and balance due update correctly without reload.",
            "priority": "High",
        },
        {
            "id": "INV-003",
            "module": "Invoice",
            "scenario": "Template switching works",
            "preconditions": "Invoice workspace is open.",
            "steps": "Switch between Minimal, Professional, and Bold Brand templates.",
            "expected": "Each template renders correctly and form values remain intact.",
            "priority": "High",
        },
        {
            "id": "INV-004",
            "module": "Invoice",
            "scenario": "Visibility toggles hide optional blocks",
            "preconditions": "Invoice workspace is open.",
            "steps": "Disable Notes and Payment Summary visibility toggles.",
            "expected": "The affected sections disappear from preview and layout remains balanced.",
            "priority": "Medium",
        },
        {
            "id": "INV-005",
            "module": "Invoice",
            "scenario": "PDF export completes",
            "preconditions": "Invoice workspace is open with valid data.",
            "steps": "Click Export PDF and open the downloaded file.",
            "expected": "An invoice PDF downloads successfully and contains the expected business, invoice number, and client text.",
            "priority": "Blocker",
        },
        {
            "id": "INV-006",
            "module": "Invoice",
            "scenario": "PNG export completes",
            "preconditions": "Invoice workspace is open with valid data.",
            "steps": "Click Export PNG and open the downloaded file.",
            "expected": "An invoice PNG downloads successfully and the document image is complete and readable.",
            "priority": "Blocker",
        },
    ]

    export_cases = [
        {
            "id": "EXP-001",
            "module": "Export",
            "scenario": "PDF files contain expected document text",
            "preconditions": "One PDF has been exported from each module.",
            "steps": "Open each downloaded PDF and review the main document identity fields.",
            "expected": "The PDF text is selectable and includes the expected company, employee, client, or voucher reference content.",
            "priority": "High",
        },
        {
            "id": "EXP-002",
            "module": "Print",
            "scenario": "Print surfaces load correctly",
            "preconditions": "A tester can allow pop-ups in the browser.",
            "steps": "Open the print flow from voucher, salary slip, and invoice workspaces.",
            "expected": "Each print page opens on the correct render-ready document surface without missing payload errors.",
            "priority": "High",
        },
        {
            "id": "EXP-003",
            "module": "Runtime",
            "scenario": "Serverless export works on deployed preview",
            "preconditions": "Testing is being performed on Vercel Preview.",
            "steps": "Run one PDF and one PNG export from each module on the deployed preview.",
            "expected": "Exports complete successfully without deployment-specific runtime failures.",
            "priority": "Blocker",
        },
    ]

    responsive_cases = [
        {
            "id": "RWD-001",
            "module": "Responsive",
            "scenario": "Desktop layout remains usable",
            "preconditions": "Use a desktop viewport.",
            "steps": "Open each module and review the form and preview layout.",
            "expected": "The workspace remains readable, preview fits properly, and controls are accessible.",
            "priority": "Medium",
        },
        {
            "id": "RWD-002",
            "module": "Responsive",
            "scenario": "Tablet/mobile layout remains usable",
            "preconditions": "Use at least one tablet or mobile-sized viewport.",
            "steps": "Open each module and review form flow, preview scaling, and action buttons.",
            "expected": "The workspace remains usable without broken layout, clipped controls, or unreadable preview content.",
            "priority": "Medium",
        },
    ]

    negative_cases = [
        {
            "id": "NEG-001",
            "module": "Validation",
            "scenario": "Required-field validation blocks export",
            "preconditions": "Open any module workspace.",
            "steps": "Clear a required field such as voucher number or equivalent core field. Attempt export.",
            "expected": "Export is blocked and a readable validation message appears.",
            "priority": "High",
        },
        {
            "id": "NEG-002",
            "module": "Render",
            "scenario": "Unavailable payload pages are understandable",
            "preconditions": "Open a print route directly without going through the workspace flow if possible.",
            "steps": "Load the print URL without a valid session payload.",
            "expected": "A clear render payload unavailable message is shown instead of a broken blank page.",
            "priority": "Medium",
        },
    ]

    for heading, cases in [
        ("Home / Navigation", home_cases),
        ("Voucher Generator", voucher_cases),
        ("Salary Slip Generator", salary_cases),
        ("Invoice Generator", invoice_cases),
        ("Export and Print", export_cases),
        ("Responsive / Layout", responsive_cases),
        ("Negative / Validation", negative_cases),
    ]:
        document.add_heading(heading, level=2)
        add_test_case_table(document, cases)

    document.add_heading("Defect Logging Notes", level=1)
    add_numbered(
        document,
        [
            "Capture at least one screenshot of the visible issue.",
            "Record the exact route and module where the issue occurred.",
            "Record the selected template name.",
            "Record the browser and viewport used.",
            "Record the input data entered before the issue occurred.",
            "For export issues, record whether the flow was Print, PDF, or PNG.",
            "Record the actual result and expected result in one concise statement each.",
        ],
    )

    document.add_heading("Sign-off Section", level=1)
    signoff = document.add_table(rows=5, cols=2)
    signoff.style = "Table Grid"
    signoff.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row, (label, value) in zip(
        signoff.rows,
        [
            ("QA Tester Name", "______________________________"),
            ("Test Date", "______________________________"),
            ("Environment", "Vercel Preview"),
            ("Result", "Pass / Fail / Blocked"),
            ("Notes", "______________________________"),
        ],
    ):
        set_cell_text(row.cells[0], label, bold=True)
        shade_cell(row.cells[0], "EADFCB")
        set_cell_text(row.cells[1], value)

    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
