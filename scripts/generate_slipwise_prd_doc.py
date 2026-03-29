from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "slipwise-rebrand-prd.docx"


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def configure_section(section) -> None:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.text = "Page "
    add_page_number(paragraph)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    for style_name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = document.styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)


def add_title_page(document: Document) -> None:
    document.add_paragraph()
    document.add_paragraph()

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Slipwise")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Atmospheric SaaS Redesign and Production Readiness PRD")
    run.bold = True
    run.font.size = Pt(16)

    document.add_paragraph()

    for label, value in [
        ("Document Version", "v2.0"),
        ("Prepared For", "Product, Design, Engineering, and QA"),
        ("Prepared On", date.today().strftime("%d %B %Y")),
        ("Primary Scope", "Phase 1 homepage and shared shell redesign"),
        ("Design Direction", "Atmospheric SaaS, white theme, big product mockup"),
    ]:
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run(value)

    document.add_page_break()


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def set_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_phase_table(document: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Phase", "Focus", "Primary Deliverable", "Acceptance Signal"]
    widths = [0.95, 1.75, 2.3, 2.4]

    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "DCEAF7")

    for phase, focus, deliverable, signal in rows:
        row = table.add_row().cells
        values = [phase, focus, deliverable, signal]
        for index, value in enumerate(values):
            row[index].width = Inches(widths[index])
            set_cell_text(row[index], value)

    document.add_paragraph()


def add_two_column_table(document: Document, left_title: str, right_title: str, rows: list[tuple[str, str]]) -> None:
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [left_title, right_title]
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_text(cell, header, bold=True)
        shade_cell(cell, "EAF2FB")

    for left, right in rows:
        row = table.add_row().cells
        set_cell_text(row[0], left)
        set_cell_text(row[1], right)

    document.add_paragraph()


def add_section_break(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    configure_section(section)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    configure_styles(document)
    configure_section(document.sections[0])
    add_title_page(document)
    configure_section(document.sections[-1])

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph(
        "Slipwise already works as a functional browser-based product for vouchers, salary slips, and invoices, but the current public-facing design does not match the quality bar expected from a deployable SaaS product. The existing homepage is too narrow, too lightweight, too utility-led, and does not create enough trust or excitement on first impression."
    )
    document.add_paragraph(
        "This PRD replaces the earlier minimal direction with a stronger creative brief: atmospheric SaaS, white-theme, product-led, and visually substantial. The product should feel modern, premium, and commercially presentable from the homepage through the shared app shell, while preserving the current stateless, serverless, Vercel-friendly architecture."
    )
    add_bullets(
        document,
        [
            "Discard the current homepage direction rather than iterating on it lightly.",
            "Make the hero larger, wider, and anchored by a real product mockup composition.",
            "Upgrade the shared shell in the same phase so the site and app feel like one product.",
            "Keep all current generator logic and export capabilities intact.",
        ],
    )

    document.add_heading("2. Product Objective", level=1)
    document.add_paragraph(
        "Slipwise should be perceived as a serious SaaS product for SMB operations, HR, admin, and finance-adjacent teams that need professional business documents fast. The redesign must increase trust, perceived quality, and conversion intent without introducing backend-heavy scope."
    )
    add_bullets(
        document,
        [
            "Position Slipwise as the clean alternative to spreadsheet formatting and ad hoc document preparation.",
            "Present vouchers, salary slips, and invoices as part of one coherent product system.",
            "Create a marketing surface strong enough for demos, QA handoff, preview deployment review, and real user evaluation.",
            "Ensure the product remains easy to deploy on Vercel without auth, persistence, or database dependencies.",
        ],
    )

    document.add_heading("3. Current State Assessment", level=1)
    document.add_heading("What Already Exists", level=2)
    add_bullets(
        document,
        [
            "Voucher, salary slip, and invoice generators are implemented.",
            "Each module has live preview, template switching, branding controls, and export paths.",
            "Print, PDF, and PNG flows already exist.",
            "The product is stateless and serverless-friendly.",
            "The codebase is already structured well enough to support a visual redesign without changing core business logic.",
        ],
    )
    document.add_heading("What Is Not Working in the Current Design", level=2)
    add_bullets(
        document,
        [
            "The homepage feels too thin and generic relative to the desired SaaS reference quality.",
            "Typography scale is too small and lacks enough visual drama in the hero.",
            "The layout width and section composition do not create a premium, product-marketing feel.",
            "The page relies on small cards and lightweight content blocks instead of stronger art direction.",
            "The current shell and homepage are not yet cohesive enough to feel like one intentional product.",
        ],
    )

    document.add_heading("4. Locked Strategic Decisions", level=1)
    add_two_column_table(
        document,
        "Decision",
        "Locked Choice",
        [
            ("Overall direction", "Atmospheric SaaS, not sparse utility design."),
            ("Theme", "White/light theme with cool neutral surfaces and strong blue accent."),
            ("Hero treatment", "Big product mockup dominates the visual story."),
            ("Phase 1 scope", "Homepage plus shared app shell, not homepage-only."),
            ("Phase 2 scope", "Generator workspace polish after the new public and shell system is established."),
            ("Architecture constraint", "Keep the product stateless, serverless, and Vercel-friendly."),
        ],
    )

    document.add_heading("5. Audience and Positioning", level=1)
    document.add_heading("Primary User", level=2)
    document.add_paragraph(
        "SMB operations, HR, and admin users who actually create and share these documents in day-to-day work."
    )
    document.add_heading("Secondary User", level=2)
    document.add_paragraph(
        "Founders, office managers, and finance leads who evaluate the product based on professionalism, speed, and ease of use."
    )
    document.add_heading("Positioning Statement", level=2)
    document.add_paragraph(
        "Slipwise is a browser-based document workflow product that helps teams create vouchers, salary slips, and invoices with a faster and cleaner workflow than spreadsheets or manually formatted documents."
    )
    document.add_heading("Messaging Priorities", level=2)
    add_bullets(
        document,
        [
            "Professional output without design work.",
            "Structured workflows for real business documents.",
            "Immediate preview and clean export.",
            "Simple enough for admins, polished enough for the business.",
        ],
    )

    document.add_heading("6. Brand and Visual Direction", level=1)
    document.add_heading("Brand Attributes", level=2)
    add_bullets(
        document,
        [
            "Modern",
            "Precise",
            "Premium",
            "Product-led",
            "Confident",
            "Calm but not bland",
        ],
    )
    document.add_heading("Brand Anti-Patterns", level=2)
    add_bullets(
        document,
        [
            "Do not make the site feel like an internal admin tool.",
            "Do not make the design too airy or under-designed.",
            "Do not rely on tiny cards and small typography to carry the page.",
            "Do not use a beige/editorial identity for the product shell.",
            "Do not over-index on generic startup copy without product proof.",
        ],
    )
    document.add_heading("Design System Direction", level=2)
    add_bullets(
        document,
        [
            "Display font: Sora.",
            "Body/UI font: Manrope.",
            "Base palette: white, cool gray, deep slate, strong product blue, selective cyan and mint support tones.",
            "Large rounded surfaces, premium borders, controlled shadows, and more deliberate section separation.",
            "Use atmospheric gradients and glows to add richness without compromising the white-theme identity.",
            "Allow selective dark surfaces for hero benefit panels or final CTA sections.",
        ],
    )

    document.add_heading("7. Delivery Roadmap", level=1)
    add_phase_table(
        document,
        [
            (
                "Phase 1",
                "Homepage and shared shell redesign",
                "New Slipwise homepage plus upgraded shared app chrome",
                "Homepage feels commercially presentable and the app shell visually matches it",
            ),
            (
                "Phase 2",
                "Generator workspace polish",
                "Voucher, salary slip, and invoice UX refinement",
                "Each generator page feels premium and cohesive under the new system",
            ),
        ],
    )

    document.add_heading("8. Phase 1 Requirements: Homepage and Shared Shell", level=1)
    document.add_paragraph(
        "Phase 1 is the active implementation scope. It must be strong enough that a reviewer, tester, or stakeholder can land on the site and immediately feel that Slipwise is a real SaaS product. This phase includes the public homepage and the shared app shell, but not deep workspace redesign inside each generator."
    )

    document.add_heading("8.1 Homepage Information Architecture", level=2)
    add_numbered(
        document,
        [
            "Top navigation with Slipwise wordmark, key links, and two clear CTAs.",
            "Hero with large headline, stronger supporting paragraph, trust chips, and a dominant product mockup composition.",
            "Trust strip or product proof row directly beneath or integrated into the hero.",
            "Feature story section with one anchor block and several substantial supporting cards.",
            "Use case section framed by real roles: HR/Admin, Operations, Finance/Accounts.",
            "Workflow section that explains the product in three clear steps.",
            "Generator showcase that introduces voucher, salary slip, and invoice as three modules in one system.",
            "FAQ section with a smaller set of stronger questions and cleaner accordions.",
            "Final CTA section with a darker or more accent-heavy close.",
            "Minimal premium footer.",
        ],
    )

    document.add_heading("8.2 Hero Requirements", level=2)
    add_bullets(
        document,
        [
            "The hero must be significantly wider and visually heavier than the current version.",
            "The main heading should read as a product headline, not a generic utility description.",
            "The supporting paragraph should be concise but strong enough to establish value in one read.",
            "CTAs must be obvious and prioritized: primary CTA to start, secondary CTA to explore generators or product sections.",
            "Trust chips should reinforce browser-based workflow, export readiness, and supported document categories.",
            "The hero must have atmospheric background treatment so the top of the page does not feel empty or sterile.",
        ],
    )

    document.add_heading("8.3 Hero Visual / Product Mockup Requirements", level=2)
    add_bullets(
        document,
        [
            "The visual should look like software marketing, not a cluster of small stat cards.",
            "Use a large composed mockup with layered panels, surfaces, and depth.",
            "Represent Slipwise as one product with multiple generators, not three unrelated mini cards.",
            "The visual should include product-shell cues such as navigation, tabs, workspace panels, form controls, preview surface, and status/export elements.",
            "The composition may be partially fictionalized for marketing clarity, but must remain faithful to the actual product structure.",
            "The mockup should remain readable on desktop and collapse intelligently on tablet/mobile.",
        ],
    )

    document.add_heading("8.4 Homepage Content Guidelines", level=2)
    add_bullets(
        document,
        [
            "Copy must sound like a launchable SaaS product, not a development update.",
            "Avoid language such as foundation, phase, shell, scaffold, or future plan on the public page.",
            "Keep statements specific and product-focused.",
            "Use role-led language where helpful, especially in use-case sections.",
            "Prioritize clarity, confidence, and trust over cleverness.",
        ],
    )

    document.add_heading("8.5 Shared Shell Requirements", level=2)
    add_bullets(
        document,
        [
            "The shared app shell must visually align with the homepage in color, spacing, typography, and surface language.",
            "Workspace headers, section cards, action bars, and preview framing must feel premium and cohesive.",
            "Shared buttons and call-to-action hierarchy must become more deliberate and polished.",
            "Outer product chrome should become SaaS-like while keeping document preview internals stable.",
            "The redesign must not change generator logic, export APIs, or data flow.",
        ],
    )

    document.add_heading("8.6 Content and Component Architecture", level=2)
    add_bullets(
        document,
        [
            "Homepage sections should be modular and reusable rather than one large inline file.",
            "Marketing content should live in structured config/data where practical.",
            "The design token layer must be semantic so it can power both public and in-app surfaces.",
            "Module cards should be unified between the homepage and shared product shell.",
        ],
    )

    add_section_break(document)

    document.add_heading("9. Phase 2 Requirements: Generator Workspace Polish", level=1)
    document.add_paragraph(
        "Phase 2 begins after Phase 1 is approved. It deepens the redesign into the actual voucher, salary slip, and invoice workspaces without changing the product scope."
    )
    add_bullets(
        document,
        [
            "Voucher workspace: cleaner field grouping, stronger amount emphasis, improved voucher-type presentation, and cleaner notes/signature hierarchy.",
            "Salary slip workspace: clearer separation of employee, payroll, attendance, bank, and notes sections.",
            "Invoice workspace: cleaner line-item editor, stronger totals emphasis, and better structure for client/business/payment blocks.",
            "Shared improvements: stronger template picker, better export state presentation, cleaner validation feedback, and improved mobile layouts.",
        ],
    )

    document.add_heading("10. Technical Constraints", level=1)
    add_bullets(
        document,
        [
            "The product must remain fully serverless and deployable on Vercel.",
            "No auth, persistence, billing, or backend-heavy product scope may be introduced as part of this redesign.",
            "Preview, print, PDF, and PNG flows must keep working.",
            "Do not redesign export route contracts or normalize document schemas unless required for visual consistency in shared surfaces.",
            "Marketing and shell redesign must not break current route structure.",
        ],
    )

    document.add_heading("11. Public Interfaces and Allowed Code Changes", level=1)
    add_bullets(
        document,
        [
            "Allowed: Slipwise brand config, homepage content config, semantic visual tokens, shared UI primitives, shared shell components, marketing sections.",
            "Allowed: metadata and product copy updates across public-facing surfaces.",
            "Allowed: redesign of shared module cards and shared navigation/footer treatments.",
            "Not allowed in this phase: new backend services, new persistence models, dashboard/account architecture, pricing systems, or workflow automation.",
        ],
    )

    document.add_heading("12. Quality Bar and Acceptance Criteria", level=1)
    document.add_heading("Homepage Acceptance", level=2)
    add_bullets(
        document,
        [
            "The homepage must feel visually comparable in confidence to modern SaaS references, not like a lightweight utility landing page.",
            "Hero typography must be larger, wider, and more commanding than the current implementation.",
            "The hero visual must read as a meaningful product composition.",
            "Section cards must feel substantial and intentionally art-directed.",
            "The page must remain clean and premium on desktop, tablet, and mobile.",
        ],
    )
    document.add_heading("Shared Shell Acceptance", level=2)
    add_bullets(
        document,
        [
            "The app shell must clearly belong to the same product as the homepage.",
            "Shared panels, preview frames, and action groups must feel more premium and cohesive.",
            "All three generator routes must still load and remain usable.",
        ],
    )
    document.add_heading("Production Readiness Acceptance", level=2)
    add_bullets(
        document,
        [
            "The redesign builds and deploys cleanly.",
            "The product remains stable on Vercel preview.",
            "Existing exports still work after visual changes.",
            "QA can review the product without obvious brand mismatch or unfinished public messaging.",
        ],
    )

    document.add_heading("13. Testing and Verification Plan", level=1)
    add_bullets(
        document,
        [
            "Run lint, unit tests, and production build after each major integration step.",
            "Run targeted Playwright checks for homepage load and route accessibility.",
            "Do a manual desktop/tablet/mobile review of the homepage and shared shell.",
            "Re-run print/PDF/PNG smoke checks after shared-shell changes.",
            "Review typography scale, spacing, and card density visually before approval.",
        ],
    )

    document.add_heading("14. Risks and Mitigations", level=1)
    add_two_column_table(
        document,
        "Risk",
        "Mitigation",
        [
            (
                "Hero becomes decorative but not product-led",
                "Lock the hero around a real product mockup composition instead of abstract visuals alone.",
            ),
            (
                "Homepage and app shell drift apart visually",
                "Implement both in the same Phase 1 branch under one token system.",
            ),
            (
                "Redesign breaks export behavior indirectly",
                "Keep business logic untouched and rerun export-focused checks after shared-shell changes.",
            ),
            (
                "Page still feels under-designed after implementation",
                "Bias toward larger type, stronger section hierarchy, wider layout, and more substantial hero composition.",
            ),
        ],
    )

    document.add_heading("15. Explicit Non-Goals", level=1)
    add_bullets(
        document,
        [
            "No authentication or accounts.",
            "No saved drafts or persistence.",
            "No pricing, billing, or subscriptions.",
            "No dashboard or admin back office.",
            "No team collaboration features.",
            "No backend architectural expansion disguised as design work.",
        ],
    )

    document.add_heading("16. Final Direction Statement", level=1)
    document.add_paragraph(
        "Slipwise should no longer look like a neat utility with a thin landing page. It should look like a serious SaaS product with a strong homepage, a richer and more intentional visual system, and a shared shell that makes the app feel commercially ready. Phase 1 must deliver that impression decisively."
    )

    document.save(OUTPUT)
    print(f"Wrote {OUTPUT}")


if __name__ == "__main__":
    main()
